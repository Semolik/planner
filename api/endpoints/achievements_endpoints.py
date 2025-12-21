import uuid
from io import BytesIO
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Cm
from fastapi import APIRouter, Depends, HTTPException, Body, Query
from starlette.responses import StreamingResponse

from api.cruds.meetings_crud import MeetingsCRUD
from api.cruds.users_crud import UsersCRUD
from api.models.user_models import User
from api.schemas.custom_achievements import (
    AchievementRead,
    AchievementCreate,
    AchievementUpdate,
)
from api.core.users_controller import current_user
from api.cruds.custom_achievements_crud import CustomAchievementsCRUD
from api.db.session import get_async_session
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

api_router = APIRouter(prefix="/achievements", tags=["achievements"])


def add_numbered_list(doc, title, items):
    if items:
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for i, item in enumerate(items, 1):
            list_para = doc.add_paragraph()
            list_para.paragraph_format.left_indent = Cm(0.5)
            list_para.paragraph_format.first_line_indent = -Cm(0.5)

            number_run = list_para.add_run(f"{i}.\t")
            number_run.font.name = "Times New Roman"
            number_run.font.size = Pt(14)

            text_run = list_para.add_run(
                f"{item['name']} ({item['date_from'].strftime('%d.%m.%Y')})"
            )
            text_run.font.name = "Times New Roman"
            text_run.font.size = Pt(14)

        doc.add_paragraph()


@api_router.get("", response_model=list[AchievementRead])
async def get_achievements_by_year(
    year: int,
    only_custom: bool = False,
    db=Depends(get_async_session),
    current_user: User = Depends(current_user),
):
    achievements = await CustomAchievementsCRUD(db).get_user_achievements_by_year(
        user_id=current_user.id,
        year=year,
        only_custom=only_custom,
    )
    return achievements


@api_router.get("/export", response_model=bytes)
async def export_achievements_by_year(
    year: int,

    kurs: int,
    group: str ,
    is_magistracy: bool,
    date: str ,
        meetings_ids: list[uuid.UUID] = Query(None),
    db=Depends(get_async_session),
    current_user: User = Depends(current_user),
):
    meetings = await MeetingsCRUD(db).get_meetings_by_ids(meetings_ids) if meetings_ids else []
    all_achievements = await CustomAchievementsCRUD(db).get_user_achievements_by_year(
        user_id=current_user.id,
        year=year,
    )
    accepted_achievements = []
    for achievement in all_achievements:
        if (
            achievement["is_custom"] or
            achievement["link"]
        ):
            continue
        accepted_achievements.append(achievement)
    doc = Document()
    # Получаем первый раздел (секцию)
    section = doc.sections[0]

    # Устанавливаем отступы
    section.top_margin = Cm(2.17)
    section.bottom_margin = Cm(1.62)
    section.left_margin = Cm(2.79)
    section.right_margin = Cm(1.31)
    # Установка стиля Normal на Times New Roman, 14pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # Заголовок
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(
        "В конкурсную комиссию\nпо назначению повышенных\nгосударственных академических\nстипендий ФГБОУ ВО\n«Алтайский государственный\nуниверситет»"
    )
    header_run.font.name = "Times New Roman"
    header_run.line_spacing = 1
    header_run.font.size = Pt(14)

    # СПРАВКА заголовок
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("СПРАВКА")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    title_run.bold = True

    nb_space = "\u00a0"
    p1 = doc.add_paragraph(
        f"Подтверждающая, что {current_user.last_name} {current_user.first_name} {current_user.patronymic}, "
        f"студент {kurs}{nb_space}курса{nb_space if not is_magistracy else ''}{' магистратуры' if is_magistracy else ''}{nb_space}гр.{nb_space}{group} института {current_user.institute.name}, действительно "
        f"является участником и активистом студенческого объединения ФГБОУ ВО «Алтайского государственного "
        f"университета» «Объединение фотографов» c {date}."
    )
    p1_format = p1.paragraph_format
    p1_format.line_spacing = 1.5
    p1_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    designs = [
        ach
        for ach in accepted_achievements
        if ach["level_of_participation"] == "Дизайнер"
    ]
    copyrights = [
        ach
        for ach in accepted_achievements
        if ach["level_of_participation"] == "Журналист"
    ]
    other = [
        ach
        for ach in accepted_achievements
        if ach["level_of_participation"] not in ["Дизайнер", "Журналист"]
    ]
    if designs:
        add_numbered_list(doc, "Является автором следующих дизайнов:", designs)
    if copyrights:
        add_numbered_list(doc, "Является автором следующих текстов:", copyrights)
    if other:
        add_numbered_list(doc, "Является автором следующих работ:", other)

    if meetings:
        items = [
            {"name": "Собрание Объединения фотографов", "date_from": meeting.date}
            for meeting in meetings
        ]
        add_numbered_list(doc, "Посетил(а) собрания объединения:", items)

    doc.add_paragraph()  # Пустая строка перед подписями

    superusers = await UsersCRUD(db).get_users(only_superusers=True)

    def set_cell_width(cell, width):
        """Устанавливает ширину ячейки в сантиметрах"""
        tcPr = cell._element.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width * 567)))  # Конвертируем см в twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    for user in superusers:
        # Создаём невидимую таблицу
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Устанавливаем ширину через функцию (60% на левую, 40% на правую)
        set_cell_width(table.rows[0].cells[0], 10.5)  # Левая колонка
        set_cell_width(table.rows[0].cells[1], 6.5)  # Правая колонка

        # Убираем границы

        tblPr = table._element.tblPr
        tblBorders = parse_xml(
            r'<w:tblBorders %s><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders>'
            % nsdecls("w")
        )
        tblPr.append(tblBorders)

        # Левая колонка: должность
        left_cell = table.rows[0].cells[0]
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.text = "Руководитель Объединения фотографов"
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Правая колонка: инициалы
        right_cell = table.rows[0].cells[1]
        right_paragraph = right_cell.paragraphs[0]
        initials = f"{user.last_name} {user.first_name[0].upper()}."
        if user.patronymic:
            initials += f" {user.patronymic[0].upper()}."
        right_paragraph.text = initials
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    filename = f"Справка {current_user.last_name} {current_user.first_name}.docx"
    quoted_filename = quote(filename)
    headers = {"Content-Disposition": f"attachment; filename*=UTF-8''{quoted_filename}"}
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )
