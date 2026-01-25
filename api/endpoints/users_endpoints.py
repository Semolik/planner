from api.core.users_controller import (
    current_superuser,
    current_user,
    optional_current_user,
)
from api.cruds.custom_achievements_crud import CustomAchievementsCRUD
from api.cruds.institutes_crud import InstitutesCRUD
from api.cruds.meetings_crud import MeetingsCRUD
from api.cruds.periods_crud import PeriodsCRUD
from api.cruds.tasks_crud import TasksCRUD
from api.cruds.users_crud import UsersCRUD
from api.db.session import get_async_session
from api.models.user_models import User, UserRole
from api.schemas.achievements import AchievementRead
from api.schemas.events import TaskRead, TypedTaskReadFull, UserReadAdmin
from api.schemas.users import UserRead, UserReadWithEmail, UserUpdate
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shared import OxmlElement
from docx.shared import Cm, Pt
from fastapi import APIRouter, Depends, HTTPException, Query
from io import BytesIO
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from pydantic import TypeAdapter
from starlette.responses import StreamingResponse
from typing import List, Literal, Union
from urllib.parse import quote
import uuid

api_router = APIRouter(prefix="/users", tags=["users"])


async def update_handler(db, user: User, user_data: UserUpdate, current_user: User):
    users_crud = UsersCRUD(db)
    if not current_user.is_superuser and current_user.id != user.id:
        raise HTTPException(status_code=403, detail="Нет доступа")
    if current_user.is_superuser:
        if user.id == current_user.id and not user_data.is_superuser:
            raise HTTPException(
                status_code=403, detail="Нельзя понизить себя до обычного пользователя"
            )
    if (
        user_data.username != user.username
        and await users_crud.get_user_by_username(user_data.username) is not None
    ):
        raise HTTPException(
            status_code=400, detail="Пользователь с таким username уже существует"
        )
    if user_data.institute_id != user.institute_id:
        institute_crud = InstitutesCRUD(db)
        if not await institute_crud.get_institute_by_id(user_data.institute_id):
            raise HTTPException(status_code=400, detail="Институт не найден")
    await users_crud.update_user(
        user=user, user_data=user_data, update_as_superuser=current_user.is_superuser
    )
    return await users_crud.get_user_by_id(user.id)


@api_router.put("/{user_id}", response_model=UserReadWithEmail)
async def update_user(
    user: UserUpdate,
    user_id: uuid.UUID,
    db=Depends(get_async_session),
    current_user: User = Depends(current_superuser),
):
    users_crud = UsersCRUD(db)
    db_user = await users_crud.get_user_by_id(user_id)
    if db_user is None:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    return await update_handler(
        db=db, user=db_user, user_data=user, current_user=current_user
    )


@api_router.get("/me", response_model=UserReadWithEmail, name="users:current_user")
async def get_user_me(
    db=Depends(get_async_session), current_user: User = Depends(current_user)
):
    return await UsersCRUD(db).get_user_by_id(current_user.id)


@api_router.get("/{user_id}", response_model=Union[UserReadWithEmail, UserRead])
async def get_user(
    user_id: uuid.UUID,
    db=Depends(get_async_session),
    current_user: User = Depends(current_user),
):
    user = await UsersCRUD(db).get_user_by_id(user_id)
    if user is None:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    if current_user.is_superuser:
        return UserReadWithEmail.model_validate(user)
    return UserRead.model_validate(user)


@api_router.get("", response_model=List[Union[UserRead, UserReadAdmin]])
async def get_users(
    search: str = None,
    page: int = Query(1, ge=1),
    order_by: Literal["last_name", "birth_date"] = "last_name",
    order: Literal["asc", "desc"] = "asc",
    superusers_to_top: bool = False,
    only_superusers: bool = False,
    filter_role: UserRole = None,
    db=Depends(get_async_session),
    current_user: User = Depends(optional_current_user),
):
    users = await UsersCRUD(db).get_users(
        order_by=order_by,
        order=order,
        search=search,
        page=page,
        superusers_to_top=superusers_to_top,
        only_superusers=only_superusers,
        filter_role=filter_role,
    )
    if current_user and current_user.is_superuser:
        return TypeAdapter(List[UserReadAdmin]).validate_python(users)
    return TypeAdapter(List[UserRead]).validate_python(users)


@api_router.delete(
    "/{user_id}", status_code=204, dependencies=[Depends(current_superuser)]
)
async def delete_user(
    user_id: uuid.UUID,
    db=Depends(get_async_session),
):
    users_crud = UsersCRUD(db)
    user = await users_crud.get_user_by_id(user_id)
    if user is None:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    await users_crud.delete(user)


@api_router.get(
    "/{user_id}/typed_tasks/completed", response_model=list[TypedTaskReadFull]
)
async def get_user_completed_typed_tasks(
    user_id: uuid.UUID,
    period_id: uuid.UUID = Query(...),
    db=Depends(get_async_session),
    current_db_user: User = Depends(current_user),
):
    user = await UsersCRUD(db).get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    if user.id != current_db_user.id and not current_db_user.is_superuser:
        raise HTTPException(status_code=403, detail="Доступ запрещен")
    periods_crud = PeriodsCRUD(db)
    period = await periods_crud.get_period_by_id(period_id)
    if not period:
        raise HTTPException(status_code=404, detail="Период не найден")
    return await TasksCRUD(db).get_user_completed_typed_tasks(
        user_id=user_id,
        period_start=period.period_start,
        period_end=period.period_end,
    )


@api_router.get("/{user_id}/birthdays-tasks", response_model=list[TaskRead])
async def get_user_birthdays_tasks(
    user_id: uuid.UUID,
    offset: int = Query(0, ge=0),
    limit: int = Query(10, ge=1),
    db=Depends(get_async_session),
    current_db_user: User = Depends(current_user),
):
    user = await UsersCRUD(db).get_user_by_id(user_id)
    if not user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    if user.id != current_db_user.id and not current_db_user.is_superuser:
        raise HTTPException(status_code=403, detail="Доступ запрещен")
    return await TasksCRUD(db).get_user_birthdays_tasks(
        user_id=user_id,
        offset=offset,
        limit=limit,
    )


def add_numbered_list(doc, title, items):
    if items:
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for i, item in enumerate(items, 1):
            list_para = doc.add_paragraph()
            list_para.paragraph_format.left_indent = Cm(0.5)
            list_para.paragraph_format.first_line_indent = -Cm(0.5)

            number_run = list_para.add_run(f"{i}.\t")
            number_run.font.name = "Times New Roman"
            number_run.font.size = Pt(14)

            text_run = list_para.add_run(
                f"{item['name']} ({item['date_from'].strftime('%d.%m.%Y')})"
            )
            text_run.font.name = "Times New Roman"
            text_run.font.size = Pt(14)

        doc.add_paragraph()


@api_router.get("/{user_id}/achievements", response_model=list[AchievementRead])
async def get_achievements_by_year(
    user_id: uuid.UUID,
    year: int,
    only_custom: bool = False,
    db=Depends(get_async_session),
    current_user: User = Depends(current_user),
):
    target_user = await UsersCRUD(db).get_user_by_id(user_id)
    if not target_user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    if current_user.id != user_id and not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Нет доступа")

    achievements = await CustomAchievementsCRUD(db).get_user_achievements_by_year(
        user_id=user_id,
        year=year,
        only_custom=only_custom,
    )
    return achievements


@api_router.get("/{user_id}/achievements/export", response_model=bytes)
async def export_achievements_by_year(
    user_id: uuid.UUID,
    year: int,
    kurs: int,
    group: str,
    is_magistracy: bool,
    date: str,
    meetings_ids: list[uuid.UUID] = Query(None),
    db=Depends(get_async_session),
    current_user: User = Depends(current_user),
):
    target_user = await UsersCRUD(db).get_user_by_id(user_id)
    if not target_user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    if current_user.id != user_id and not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Нет доступа")

    meetings = (
        await MeetingsCRUD(db).get_meetings_by_ids(meetings_ids) if meetings_ids else []
    )
    all_achievements = await CustomAchievementsCRUD(db).get_user_achievements_by_year(
        user_id=user_id,
        year=year,
    )
    accepted_achievements = []
    for achievement in all_achievements:
        if achievement["is_custom"] or achievement["link"]:
            continue
        accepted_achievements.append(achievement)
    doc = Document()
    # Получаем первый раздел (секцию)
    section = doc.sections[0]

    # Устанавливаем отступы
    section.top_margin = Cm(2.17)
    section.bottom_margin = Cm(1.62)
    section.left_margin = Cm(2.79)
    section.right_margin = Cm(1.31)
    # Установка стиля Normal на Times New Roman, 14pt
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)

    # Заголовок
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(
        "В конкурсную комиссию\nпо назначению повышенных\nгосударственных академических\nстипендий ФГБОУ ВО\n«Алтайский государственный\nуниверситет»"
    )
    header_run.font.name = "Times New Roman"
    header_run.line_spacing = 1
    header_run.font.size = Pt(14)

    # СПРАВКА заголовок
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("СПРАВКА")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(14)
    title_run.bold = True

    nb_space = "\u00a0"
    p1 = doc.add_paragraph(
        f"Подтверждающая, что {target_user.last_name} {target_user.first_name} {target_user.patronymic}, "
        f"студент {kurs}{nb_space}курса{nb_space if not is_magistracy else ''}{' магистратуры' if is_magistracy else ''}{nb_space}гр.{nb_space}{group} института {target_user.institute.name}, действительно "
        f"является участником и активистом студенческого объединения ФГБОУ ВО «Алтайского государственного "
        f"университета» «Объединение фотографов» c {date}."
    )
    p1_format = p1.paragraph_format
    p1_format.line_spacing = 1.5
    p1_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    designs = [
        ach
        for ach in accepted_achievements
        if ach["level_of_participation"] == "Дизайнер"
    ]
    copyrights = [
        ach
        for ach in accepted_achievements
        if ach["level_of_participation"] == "Журналист"
    ]
    other = [
        ach
        for ach in accepted_achievements
        if ach["level_of_participation"] not in ["Дизайнер", "Журналист"]
    ]
    if designs:
        add_numbered_list(doc, "Является автором следующих дизайнов:", designs)
    if copyrights:
        add_numbered_list(doc, "Является автором следующих текстов:", copyrights)
    if other:
        add_numbered_list(doc, "Является автором следующих работ:", other)

    if meetings:
        items = [
            {"name": "Собрание Объединения фотографов", "date_from": meeting.date}
            for meeting in meetings
        ]
        add_numbered_list(doc, "Посетил(а) собрания объединения:", items)

    doc.add_paragraph()  # Пустая строка перед подписями

    superusers = await UsersCRUD(db).get_users(only_superusers=True)

    def set_cell_width(cell, width):
        """Устанавливает ширину ячейки в сантиметрах"""
        tcPr = cell._element.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width * 567)))  # Конвертируем см в twips
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)

    for user in superusers:
        # Создаём невидимую таблицу
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Устанавливаем ширину через функцию (60% на левую, 40% на правую)
        set_cell_width(table.rows[0].cells[0], 10.5)  # Левая колонка
        set_cell_width(table.rows[0].cells[1], 6.5)  # Правая колонка

        # Убираем границы

        tblPr = table._element.tblPr
        tblBorders = parse_xml(
            r'<w:tblBorders %s><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders>'
            % nsdecls("w")
        )
        tblPr.append(tblBorders)

        # Левая колонка: должность
        left_cell = table.rows[0].cells[0]
        left_paragraph = left_cell.paragraphs[0]
        left_paragraph.text = "Руководитель Объединения фотографов"
        left_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Правая колонка: инициалы
        right_cell = table.rows[0].cells[1]
        right_paragraph = right_cell.paragraphs[0]
        initials = f"{user.last_name} {user.first_name[0].upper()}."
        if user.patronymic:
            initials += f" {user.patronymic[0].upper()}."
        right_paragraph.text = initials
        right_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    filename = f"Справка {target_user.last_name} {target_user.first_name}.docx"
    quoted_filename = quote(filename)
    headers = {"Content-Disposition": f"attachment; filename*=UTF-8''{quoted_filename}"}
    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@api_router.get("/{user_id}/achievements/export-excel")
async def export_achievements_excel(
    user_id: uuid.UUID,
    year: int,
    participant_date: date,
    participant_link: str,
    meetings_ids: List[uuid.UUID] = Query(None),
    db=Depends(get_async_session),
    current_user: User = Depends(current_user),
):
    target_user = await UsersCRUD(db).get_user_by_id(user_id)
    if not target_user:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    if current_user.id != user_id and not current_user.is_superuser:
        raise HTTPException(status_code=403, detail="Нет доступа")

    # Получаем планерки
    meetings = (
        await MeetingsCRUD(db).get_meetings_by_ids(meetings_ids) if meetings_ids else []
    )
    meetings_count = len(meetings)

    # Получаем все достижения за год
    all_achievements = await CustomAchievementsCRUD(db).get_user_achievements_by_year(
        user_id=user_id,
        year=year,
    )

    # ФИЛЬТРАЦИЯ: убираем is_custom == false где "Журналист"
    filtered_achievements = []
    journalist_posts_count = 0

    for achievement in all_achievements:
        if (
            not achievement.get("is_custom")
            and achievement.get("level_of_participation") == "Журналист"
        ):
            journalist_posts_count += 1
        else:
            filtered_achievements.append(achievement)

    filtered_achievements.insert(
        0,
        {
            "name": "Участник студенческого объединения ФГБОУ ВО «Алтайского государственного университета» «Объединение фотографов»",
            "date_from": participant_date,
            "level_of_participation": "Участник",
            "achievement_level": "университетский",
            "link": participant_link,
            "score": 10,
            "is_participant": True,
        },
    )

    # === АГРЕГАЦИЯ ПОСТОВ ЖУРНАЛИСТА ✅ ИСПРАВЛЕННЫЕ ЦИФРЫ ===
    if journalist_posts_count >= 5:
        if journalist_posts_count <= 10:
            posts_level = "5-10"
            posts_score = 20
        elif journalist_posts_count <= 20:
            posts_level = "11-20"
            posts_score = 80
        else:
            posts_level = "более 21"
            posts_score = 150

        filtered_achievements.append(
            {
                "name": "Посты в социальных сетях студенческого объединения ФГБОУ ВО «Алтайского государственного университета» «Объединения фотографов»",
                "date_from": None,
                "level_of_participation": "Журналист",
                "achievement_level": posts_level,
                "link": "",
                "score": posts_score,
                "is_social_posts": True,
            }
        )

    # === АГРЕГАЦИЯ ПЛАНЕРОК (оставляем как было) ===
    if meetings_count >= 8:
        if meetings_count <= 12:
            meetings_level = "8-12"
            meetings_score = 30
        elif meetings_count <= 23:
            meetings_level = "13-23"
            meetings_score = 50
        else:
            meetings_level = "более 24"
            meetings_score = 80

        filtered_achievements.append(
            {
                "name": "Посещение собраний студенческого объединения ФГБОУ ВО «Алтайского государственного университета» «Объединения фотографов»",
                "date_from": None,
                "level_of_participation": "Участник",
                "achievement_level": meetings_level,
                "link": participant_link,
                "score": meetings_score,
                "is_meetings": True,
            }
        )

    # Сортируем: участник первый, потом по дате, агрегации в конце
    filtered_achievements.sort(
        key=lambda x: (
            x.get("is_participant", False),
            x.get("is_social_posts", False) or x.get("is_meetings", False),
            -(x.get("date_from") or date.min).year if x.get("date_from") else 9999,
        )
    )

    # Создаем workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Достижения"

    # Заголовки
    headers = [
        "Номер",
        "Название мероприятия",
        "Дата проведения",
        "Уровень участия",
        "Уровень мероприятия",
        "Ссылка на подтверждение",
        "Балл",
    ]

    # Стили - ПЕРЕНОС СТРОК ДЛЯ ВСЕХ ЯЧЕЕК
    header_font = Font(name="Times New Roman", size=12, bold=True)
    header_fill = PatternFill(
        start_color="CCCCCC", end_color="CCCCCC", fill_type="solid"
    )
    center_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_align = Alignment(horizontal="left", vertical="top", wrap_text=True)

    hyperlink_font = Font(
        name="Times New Roman", size=11, color="0000FF", underline="single"
    )

    # Заголовки
    for col_num, header_text in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_num, value=header_text)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_align

    # Заполняем данными
    for row_num, achievement in enumerate(filtered_achievements, 2):
        # Номер
        ws.cell(row=row_num, column=1, value=row_num - 1).alignment = center_align

        # Название
        name_cell = ws.cell(row=row_num, column=2, value=achievement.get("name", ""))
        name_cell.alignment = left_align

        # Дата
        date_cell = ws.cell(row=row_num, column=3)
        date_value = achievement.get("date_from")
        if isinstance(date_value, date):
            date_cell.value = date_value
            date_cell.number_format = "DD.MM.YYYY"
        date_cell.alignment = center_align

        # Уровень участия
        ws.cell(
            row=row_num, column=4, value=achievement.get("level_of_participation", "")
        ).alignment = left_align

        # Уровень мероприятия
        ws.cell(
            row=row_num, column=5, value=achievement.get("achievement_level", "")
        ).alignment = left_align

        # Ссылка
        link = achievement.get("link", "")
        link_cell = ws.cell(row=row_num, column=6)
        if link:
            link_cell.value = link
            link_cell.alignment = left_align
            if link.lower().startswith(("http://", "https://")):
                link_cell.hyperlink = link
                link_cell.font = hyperlink_font
            else:
                link_cell.font = Font(name="Times New Roman", size=11)
        else:
            link_cell.value = ""

        # Балл
        score_cell = ws.cell(row=row_num, column=7, value=achievement.get("score", 0))
        score_cell.alignment = center_align
        score_cell.number_format = "0"

    # Ширины колонок
    column_widths = {"A": 8, "B": 55, "C": 14, "D": 20, "E": 20, "F": 45, "G": 8}
    for col_letter, width in column_widths.items():
        ws.column_dimensions[col_letter].width = width

    # Автоматическая высота для ВСЕХ строк
    ws.row_dimensions[1].height = 25
    for row in range(2, len(filtered_achievements) + 2):
        max_lines = 1
        for col in range(1, 8):
            text = str(ws.cell(row=row, column=col).value or "")
            lines = len(text) // 45 + 1
            max_lines = max(max_lines, lines)
        ws.row_dimensions[row].height = max(15, 15 * max_lines + 5)

    # Рамки
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )
    for row in ws.iter_rows():
        for cell in row:
            cell.border = thin_border

    # Сохранение
    file_stream = BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)

    filename = f"Выгрузка_достижения_{year}_{target_user.last_name}_{target_user.first_name}.xlsx"
    quoted_filename = quote(filename)

    headers = {"Content-Disposition": f"attachment; filename*=UTF-8''{quoted_filename}"}

    return StreamingResponse(
        file_stream,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers=headers,
    )
